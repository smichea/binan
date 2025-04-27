import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

# Function to convert Markdown file to Word document with footnotes and bold formatting
def markdown_to_docx(md_filename, docx_filename):
    # Read markdown content
    with open(md_filename, 'r', encoding='utf-8') as file:
        content = file.read()

    # Find markdown links
    link_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    links = link_pattern.findall(content)

    footnotes = []

    # Replace markdown links with footnote indicators
    def replace_links(match):
        footnotes.append(match.group(2))
        return f"{match.group(1)}[{len(footnotes)}]"

    content_with_footnotes = link_pattern.sub(replace_links, content)

    # Create Word document
    doc = Document()

    # Add content by paragraphs
    paragraphs = content_with_footnotes.split('\n\n')
    for para in paragraphs:
        para = para.strip()
        if para.startswith('# '):
            doc.add_heading(para[2:], level=1)
        elif para.startswith('## '):
            doc.add_heading(para[3:], level=2)
        elif para.startswith('### '):
            doc.add_heading(para[4:], level=3)
        else:
            # Process bold formatting
            para_parts = re.split(r'(\*\*[^*]+\*\*)', para)
            doc_para = doc.add_paragraph()
            for part in para_parts:
                if part.startswith('**') and part.endswith('**'):
                    doc_para.add_run(part[2:-2]).bold = True
                else:
                    doc_para.add_run(part)

    # Add footnotes section
    if footnotes:
        doc.add_page_break()
        doc.add_heading('References', level=2)
        for idx, footnote in enumerate(footnotes, start=1):
            doc.add_paragraph(f'[{idx}] {footnote}')

    # Save document
    doc.save(docx_filename)

# Usage example
md_filename = 'raw_files/history.md'
docx_filename = 'word/history.docx'

markdown_to_docx(md_filename, docx_filename)